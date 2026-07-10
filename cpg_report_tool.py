"""
cpg_report_tool.py

Single consolidated CodedTool that runs the ENTIRE CPG market intelligence
pipeline in plain Python -- no agent-to-agent LLM hops in between -- and
produces a Word document (.docx) with all findings and the dashboard image
embedded, plus a short 3-line summary for the chat reply.

WHY THIS EXISTS:
Earlier versions of this network chained 6-7 separate LLM agents together
(brand discovery -> market trend -> feedback -> competitive -> portfolio
comparison -> dashboard -> reporting), each one a full LLM round trip on
top of slow external API calls (Google Trends, web search, news search).
That many sequential LLM hops is what was causing timeouts. This tool
collapses ALL of that into ONE coded-tool call:
    - fetches sales/trend, feedback, and competitor data (in PARALLEL via
      a thread pool, since they're independent network calls)
    - scores sentiment
    - renders the dashboard image
    - writes a .docx report with all sections + the embedded image
    - computes a grounded 3-line summary directly from the data (no LLM
      needed to write it)

The calling agent only needs ONE LLM turn: call this tool, then relay its
"summary_3_lines" and "report_url" verbatim. That's it.

Setup:
    pip install python-docx pytrends duckduckgo-search vaderSentiment requests matplotlib numpy
    (if duckduckgo-search is unavailable, it was renamed: pip install ddgs)
    Set NEWSAPI_KEY (free at https://newsapi.org/register) for competitor news.
    Optionally set REPORTS_DIR and REPORTS_BASE_URL (see bottom of file).
"""

import base64
import io
import os
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from typing import Any, Dict, List, Union

from docx import Document
from docx.shared import Inches

from sales_data_tool import SalesDataTool
from feedback_data_tool import FeedbackDataTool
from competitor_data_tool import CompetitorDataTool
from sentiment_tool import SentimentTool
from dashboard_generator_tool import DashboardGeneratorTool

REPORTS_DIR = os.environ.get(r"C:\Users\Guest1\Desktop\training\neuro-san-studio\coded_tools\basic", "generated_reports")
REPORTS_BASE_URL = os.environ.get("https://localhost:8080", "").rstrip("/")
from neuro_san.interfaces.coded_tool import CodedTool

class CPGReportTool(CodedTool):
    """
    CodedTool that runs the full pipeline synchronously in Python and
    returns a short summary + a link to the generated .docx report.
    """

    def invoke(self, args: Dict[str, Any], sly_data: Dict[str, Any]) -> Dict[str, Any]:
        product = args.get("product")
        country = args.get("country", "India")
        time_range = args.get("time_range", "last 12 months")
        competitors: List[str] = args.get("competitors") or args.get("brands") or []

        if not product:
            return {"error": "Missing required 'product' (e.g. 'oat milk')."}
        if not competitors:
            return {"error": "Missing required 'competitors' (list of brand names to analyze)."}

        errors: List[str] = []

        # --- 1. Fetch sales/trend, feedback, and competitor data IN PARALLEL ---
        # These three are independent network calls -- running them
        # concurrently instead of one-after-another is the single biggest
        # lever for avoiding timeouts.
        with ThreadPoolExecutor(max_workers=3) as pool:
            f_sales = pool.submit(
                SalesDataTool().invoke,
                {"brands": competitors, "country": country, "time_range": time_range},
                sly_data,
            )
            f_feedback = pool.submit(
                FeedbackDataTool().invoke,
                {"brands": competitors, "limit_per_brand": 20},
                sly_data,
            )
            f_competitor = pool.submit(
                CompetitorDataTool().invoke,
                {"brands": competitors, "page_size": 10},
                sly_data,
            )

            sales_result = self._safe_result(f_sales, errors, "market trend data")
            feedback_result = self._safe_result(f_feedback, errors, "feedback data")
            competitor_result = self._safe_result(f_competitor, errors, "competitor news data")

        # --- 2. Sentiment scoring (depends on feedback text just fetched) ---
        sentiment_result = {}
        try:
            sentiment_result = SentimentTool().invoke({}, sly_data)
        except Exception as e:
            errors.append(f"sentiment scoring: {str(e)}")

        # --- 3. Dashboard image (depends on everything above, all via sly_data) ---
        dashboard_result = {}
        try:
            dashboard_result = DashboardGeneratorTool().invoke({"brands": competitors}, sly_data)
        except Exception as e:
            errors.append(f"dashboard rendering: {str(e)}")

        # --- 4. Build the .docx report ---
        try:
            report_path, filename = self._build_docx(
                product=product,
                country=country,
                time_range=time_range,
                competitors=competitors,
                sly_data=sly_data,
                errors=errors,
            )
        except Exception as e:
            return {"error": f"Failed to build report document: {str(e)}"}

        # --- 5. Grounded 3-line summary for the chat reply ---
        summary_3_lines = self._build_3_line_summary(competitors, sly_data)

        if REPORTS_BASE_URL:
            report_url = f"{REPORTS_BASE_URL}/{filename}"
        else:
            report_url = f"file://{os.path.abspath(report_path)}"

        result = {
            "summary_3_lines": summary_3_lines,
            "report_url": report_url,
            "filename": filename,
        }
        if errors:
            result["data_gaps"] = errors
        return result

    # -----------------------------------------------------------------
    @staticmethod
    def _safe_result(future, errors: List[str], label: str) -> Dict[str, Any]:
        try:
            res = future.result(timeout=45)
        except Exception as e:
            errors.append(f"{label}: {str(e)}")
            return {}
        if isinstance(res, dict) and res.get("error"):
            errors.append(f"{label}: {res['error']}")
        return res or {}

    @staticmethod
    def _build_3_line_summary(brands: List[str], sly_data: Dict[str, Any]) -> str:
        sales_per_brand = sly_data.get("last_sales_data", {}).get("per_brand", {})
        sentiment_per_brand = sly_data.get("last_sentiment_by_brand", {})
        press_activity = sly_data.get("last_press_activity", {})

        lines = []

        avg_interest = {
            b: (sales_per_brand.get(b) or {}).get("avg_interest")
            for b in brands
            if (sales_per_brand.get(b) or {}).get("avg_interest") is not None
        }
        if avg_interest:
            leader = max(avg_interest, key=avg_interest.get)
            lines.append(f"{leader} leads on search/market interest in the selected market.")
        else:
            lines.append("Market interest data was limited this run -- see the report for details.")

        pos_scores = {
            b: (sentiment_per_brand.get(b) or {}).get("proportions", {}).get("positive")
            for b in brands
            if (sentiment_per_brand.get(b) or {}).get("proportions", {}).get("positive") is not None
        }
        if pos_scores:
            best = max(pos_scores, key=pos_scores.get)
            lines.append(f"{best} has the strongest customer sentiment ({pos_scores[best]*100:.0f}% positive).")
        else:
            lines.append("Sentiment data was limited this run -- see the report for details.")

        if press_activity and any(press_activity.values()):
            top_press = max(press_activity, key=press_activity.get)
            lines.append(f"{top_press} shows the most recent press/competitive activity.")
        else:
            lines.append("No notable recent press activity was found for the tracked brands.")

        return "\n".join(lines)

    def _build_docx(
        self,
        product: str,
        country: str,
        time_range: str,
        competitors: List[str],
        sly_data: Dict[str, Any],
        errors: List[str],
    ) -> tuple[str, str]:
        doc = Document()

        doc.add_heading("CPG Market Intelligence Report", level=0)
        meta = doc.add_paragraph()
        meta.add_run(f"Product: {product}    ").bold = True
        meta.add_run(f"Country: {country}    ").bold = True
        meta.add_run(f"Time range: {time_range}").bold = True
        doc.add_paragraph(f"Brands analyzed: {', '.join(competitors)}")
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC', time.gmtime())}")

        if errors:
            doc.add_heading("Data Gaps", level=1)
            for e in errors:
                doc.add_paragraph(f"- {e}", style="List Bullet")

        # --- Dashboard image ---
        image_b64 = sly_data.get("last_dashboard_image_b64")
        dashboard_description = sly_data.get("last_dashboard_description", "")
        doc.add_heading("Portfolio Dashboard", level=1)
        if image_b64:
            image_stream = io.BytesIO(base64.b64decode(image_b64))
            doc.add_picture(image_stream, width=Inches(6.2))
            doc.add_paragraph(dashboard_description)
        else:
            doc.add_paragraph("Dashboard image unavailable for this run.")

        # --- Market trends ---
        doc.add_heading("Market Trends", level=1)
        sales_per_brand = sly_data.get("last_sales_data", {}).get("per_brand", {})
        if sales_per_brand:
            for brand in competitors:
                data = sales_per_brand.get(brand)
                if not data:
                    doc.add_paragraph(f"{brand}: no trend data available.", style="List Bullet")
                    continue
                doc.add_paragraph(
                    f"{brand}: avg. interest {data.get('avg_interest', 'N/A')}/100, "
                    f"trend {data.get('trend_direction', 'unknown')}.",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("No market trend data available for this run.")

        # --- Customer feedback / sentiment ---
        doc.add_heading("Customer Feedback", level=1)
        sentiment_per_brand = sly_data.get("last_sentiment_by_brand", {})
        share_of_voice = sly_data.get("last_share_of_voice", {})
        if sentiment_per_brand:
            for brand in competitors:
                s = sentiment_per_brand.get(brand, {})
                props = s.get("proportions", {})
                sov = share_of_voice.get(brand, 0)
                if not props:
                    doc.add_paragraph(f"{brand}: no feedback data available.", style="List Bullet")
                    continue
                doc.add_paragraph(
                    f"{brand}: {props.get('positive', 0)*100:.0f}% positive / "
                    f"{props.get('neutral', 0)*100:.0f}% neutral / "
                    f"{props.get('negative', 0)*100:.0f}% negative "
                    f"(share of voice: {sov*100:.0f}%).",
                    style="List Bullet",
                )
            doc.add_paragraph(
                "Note: feedback is drawn from web-search review/complaint snippets, "
                "a directional signal rather than a curated review database."
            )
        else:
            doc.add_paragraph("No feedback data available for this run.")

        # --- Competitive landscape ---
        doc.add_heading("Competitive / Press Activity", level=1)
        articles_by_brand = sly_data.get("last_competitor_articles_by_brand") or {}
        press_activity = sly_data.get("last_press_activity", {})
        if press_activity:
            for brand in competitors:
                count = press_activity.get(brand, 0)
                doc.add_paragraph(f"{brand}: {count} recent news articles found.", style="List Bullet")
                for article in (articles_by_brand.get(brand) or [])[:3]:
                    title = article.get("title") or "(untitled)"
                    source = article.get("source") or "unknown source"
                    doc.add_paragraph(f"    - {title} ({source})", style="List Bullet 2")
        else:
            doc.add_paragraph("No competitive/press data available for this run.")

        os.makedirs(REPORTS_DIR, exist_ok=True)
        filename = f"cpg_report_{self._slugify(product)}_{int(time.time())}_{uuid.uuid4().hex[:6]}.docx"
        path = os.path.join(REPORTS_DIR, filename)
        doc.save(path)
        return path, filename

    @staticmethod
    def _slugify(text: str) -> str:
        return "".join(c if c.isalnum() else "_" for c in text.lower())[:40]


# ---------------------------------------------------------------------
# HOSTING NOTE -- how to get a real http:// link instead of file://
# ---------------------------------------------------------------------
# This tool always saves the .docx to REPORTS_DIR on disk and returns a
# file:// path by default (works for local testing -- double-click opens
# it in Word). For a real shareable http:// link, serve REPORTS_DIR over
# HTTP and set REPORTS_BASE_URL to match, e.g.:
#
#   Mount on your existing neuro-san-studio FastAPI server:
#       from fastapi.staticfiles import StaticFiles
#       app.mount("/reports", StaticFiles(directory="generated_reports"))
#   Then: REPORTS_DIR=generated_reports
#         REPORTS_BASE_URL=http://localhost:8080/reports
#
#   OR run a standalone static server:
#       python -m http.server 8000 --directory generated_reports
#   Then: REPORTS_BASE_URL=http://localhost:8000

    async def async_invoke(self, args: Dict[str, Any], sly_data: Dict[str, Any]) -> Union[Dict[str, Any], str]:
            """
            Delegates to the synchronous invoke method because it's quick, non-blocking.
            """
            return self.invoke(args, sly_data)